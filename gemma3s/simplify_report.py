import os
import io
import base64
import torch
import nibabel as nib
import numpy as np
import cv2
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import AutoProcessor, AutoModelForImageTextToText


class MedGemmaSimplify:

    def __init__(self, hf_token=None, device=None):

        if hf_token:
            from huggingface_hub import login
            login(token=hf_token)

        self.device = device if device else ("cuda" if torch.cuda.is_available() else "cpu")

        self.processor = AutoProcessor.from_pretrained("google/medgemma-1.5-4b-it")
        self.model = AutoModelForImageTextToText.from_pretrained(
            "google/medgemma-1.5-4b-it",
            torch_dtype=torch.bfloat16,
        ).to(self.device)

    # UTILITIES

    @staticmethod
    def normalize_slice(slice_array):
        p99_5 = np.percentile(slice_array, 99.5)
        p0_5 = np.percentile(slice_array, 0.5)
        norm = (slice_array - p0_5) / (p99_5 - p0_5)
        norm = np.clip(norm, 0, 1)
        return (norm * 255).astype(np.uint8)

    @staticmethod
    def bbox_from_mask(mask):
        ys, xs = np.where(mask > 0)
        if len(xs) == 0 or len(ys) == 0:
            return None
        return xs.min(), ys.min(), xs.max(), ys.max()

    @staticmethod
    def draw_bbox(image_uint8, bbox):
        img_rgb = cv2.cvtColor(image_uint8, cv2.COLOR_GRAY2RGB)
        x_min, y_min, x_max, y_max = bbox
        cv2.rectangle(img_rgb, (x_min, y_min), (x_max, y_max), (0, 255, 0), 2)
        return img_rgb

    @staticmethod
    def encode_image(img):
        with io.BytesIO() as f:
            from PIL import Image
            Image.fromarray(img).save(f, format="JPEG")
            return "data:image/jpeg;base64," + base64.b64encode(f.getvalue()).decode()

    @staticmethod
    def save_docx(text, filename, title):
        document = Document()
        title_para = document.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(
            f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        document.add_paragraph("")
        for line in text.split("\n"):
            document.add_paragraph(line.strip())
        document.save(filename)

    # MAIN PIPELINE

    def generate_reports(self, mr_volume_path, slice_index, segmentation_mask_path, manual_report_context, 
                         clinical_docx="Clinical_MRI_Report.docx", patient_docx="Patient_Friendly_Report.docx"):

        # Load Data 
        mri_vol = nib.load(mr_volume_path).get_fdata()
        seg_vol = nib.load(segmentation_mask_path).get_fdata()

        raw_slice = mri_vol[:, :, slice_index]
        seg_slice = seg_vol[:, :, slice_index]

        norm_slice = self.normalize_slice(raw_slice)
        bbox = self.bbox_from_mask(seg_slice)

        if bbox is None:
            raise ValueError("Segmentation mask is empty for selected slice.")

        bbox_img = self.draw_bbox(norm_slice, bbox)

        # CLINICAL REPORT

        instruction = (
"You are a senior neuroradiologist reporting assistant preparing a finalized imaging report.\n\n"

"You are shown an axial flair brain MRI image. "
"A rectangular bounding box highlights a region of interest.\n\n"

"You are also provided with **background anatomical context** derived from "
"prior volumetric analysis of this subject. This context may help orient "
"anatomical descriptions.\n\n"

"IMPORTANT OUTPUT RULES (Follow strictly!!):\n"
"- Do NOT include your reasoning, thoughts, analysis, or checklists\n"
"- Do NOT describe your planning process\n"
"- Do NOT include confidence scores or self-evaluation\n"
"- Output ONLY the final imaging report sections requested.\n\n"

"Describe the visual characteristics of the tissue inside the bounded region "
"based on what is visible in the image.\n\n"
)

        query_text = (
"Generate a formal neuroradiology-style clinical brain MRI report based on the provided image.\n\n"

"The report must strictly follow authentic real-world hospital radiology formatting standards.\n\n"

"STRUCTURE THE REPORT AS FOLLOWS:\n\n"

"SEQUENCES:\n"
"- Briefly state the MRI sequence(s) visible in the provided image.\n\n"

"FINDINGS:\n"
"- Divide this section into 3-4 well-structured paragraphs.\n"
"- Describe the imaging observations in natural, well-written paragraphs form.\n"
"- Include lesion characterstics, enhancement pattern (if visible), location, "
"extent and any associated mass effect or edema.\n"
"- Include anatomical extent and lesion size.\n"
"- Specify approximate anatomical location, referencing the provided context if helpful\n"
"- Comment on visible mass effect or surrounding changes, if present\n\n"
"- Use professional radiology prose, not bullet points.\n"

"IMPRESSION:\n"
"- Provide 3-4 concise, clinically styled numbered statements, summarizing the dominant findings\n"
"- Synthesize the findings rather than repeating them verbatim.\n"
"- Avoid redundancy, do not repeat the findings.\n\n"

"\nTYPICAL NEXT STEPS:\n"
"- Write in a separate single paragraph.\n"
"- Outline common follow-up imaging or correlation steps used in practice and treatment protocol\n"
"- You can also suggest some steps which should be followed in such clinical scenarios."
"- Keep this informational, not prescriptive\n\n"

"Do not include reasoning, explanations, confidence scores, or educational commentary.\n"
"Do not merge different sections in report."
"Do not restate the prompt.\n"
"Write in authentic radiology report tone."
)

        content = [
            {"type": "text", "text": instruction},
            {"type": "text", "text": manual_report_context},
            {"type": "image", "image": self.encode_image(bbox_img)},
            {"type": "text", "text": query_text},
        ]

        messages = [{"role": "user", "content": content}]

        inputs = self.processor.apply_chat_template(
            messages,
            add_generation_prompt=True,
            tokenize=True,
            return_tensors="pt",
            return_dict=True,
        )

        inputs = {
            "input_ids": inputs["input_ids"].to(self.device),
            "attention_mask": inputs["attention_mask"].to(self.device),
            "pixel_values": inputs["pixel_values"].to(self.device, dtype=torch.bfloat16),
        }

        with torch.inference_mode():
            output = self.model.generate(
                **inputs,
                max_new_tokens=2000,
                do_sample=False,
                use_cache=True,
            )

        full_text = self.processor.post_process_image_text_to_text(
            output, skip_special_tokens=True
        )[0]

        prompt_text = self.processor.post_process_image_text_to_text(
            inputs["input_ids"], skip_special_tokens=True
        )[0]

        clinical_report = (
            full_text[len(prompt_text):].strip()
            if full_text.startswith(prompt_text)
            else full_text.strip()
        )

        self.save_docx(clinical_report, clinical_docx, "Neuroradiology MRI Report")

        # PATIENT FRIENDLY REPORT

        patient_instruction = (
"You are a medical communication assistant.\n\n"

"You will be given a formal clinical MRI report written for healthcare professionals.\n\n"

"Your task is to rewrite the report in clear, simple language suitable for a general patient.\n\n"

"RULES:\n"
"- Do NOT use complex medical jargon unless you explain it.\n"
"- Explain key imaging terms in simple language.\n"
"- Keep sentences short and clear.\n"
"- Avoid giving medical advice.\n"
"- Do not speculate beyond what is written in the report.\n\n"

"STRUCTURE YOUR RESPONSE AS:\n\n"

"## What Was Seen on the MRI\n"
"- Clear explanation in plain language.\n\n"

"## What This Means in General\n"
"- High-level explanation of what these findings usually indicate.\n\n"

"## Possible Implications\n"
"- 3–5 simple bullet points describing what this could mean.\n\n"

"Keep the tone calm, reassuring, and informative.\n"
"Do not include reasoning or analysis steps."
)

        patient_content = [
            {"type": "text", "text": patient_instruction},
            {"type": "text", "text": "FORMAL MRI REPORT:\n\n" + clinical_report},
        ]

        patient_messages = [{"role": "user", "content": patient_content}]

        inputs = self.processor.apply_chat_template(
            patient_messages,
            add_generation_prompt=True,
            tokenize=True,
            return_tensors="pt",
            return_dict=True,
        )

        inputs = {
            "input_ids": inputs["input_ids"].to(self.device),
            "attention_mask": inputs["attention_mask"].to(self.device),
        }

        with torch.inference_mode():
            output = self.model.generate(
                **inputs,
                max_new_tokens=1000,
                do_sample=False,
                use_cache=True,
            )

        full_text = self.processor.post_process_image_text_to_text(
            output, skip_special_tokens=True
        )[0]

        prompt_text = self.processor.post_process_image_text_to_text(
            inputs["input_ids"], skip_special_tokens=True
        )[0]

        patient_report = (
            full_text[len(prompt_text):].strip()
            if full_text.startswith(prompt_text)
            else full_text.strip()
        )

        self.save_docx(patient_report, patient_docx, "Patient-Friendly MRI Explanation")

        return clinical_report, patient_report
